from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── helpers ──

def set_font(run, size=10, bold=False, color=None, name='Malgun Gothic'):
    run.font.name = name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_font(run, name='Malgun Gothic')
    return h

def add_para(text, bold=False, size=10, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    if align:
        p.alignment = align
    return p

def add_box(text, bg='F5F5F5'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=9)
    p.paragraph_format.left_indent = Cm(0.8)
    pr = p._element.get_or_add_pPr()
    pr.append(pr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): bg
    }))
    return p

def add_bullet(text, bold_prefix=None, color=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True, color=color)
        r2 = p.add_run(text)
        set_font(r2, color=color)
    else:
        r = p.add_run(text)
        set_font(r, color=color)

def cell_bg(cell, fill):
    pr = cell._element.get_or_add_tcPr()
    pr.append(pr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): fill
    }))

def add_table(headers, rows, col_widths=None, header_bg='F5F0EB'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        set_font(r, size=9, bold=True)
        cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            set_font(r, size=9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def add_color_table(colors):
    t = doc.add_table(rows=1+len(colors), cols=4)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['컬러', '이름', 'HEX', '용도']):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        set_font(r, size=9, bold=True)
        cell_bg(c, 'F5F0EB')
    for ri, cl in enumerate(colors):
        row = t.rows[ri+1]
        cell_bg(row.cells[0], cl['hex'].replace('#', ''))
        for ci, val in enumerate([cl['name'], cl['hex'], cl['usage']], 1):
            c = row.cells[ci]
            c.text = ''
            r = c.paragraphs[0].add_run(val)
            set_font(r, size=9)
    doc.add_paragraph()

# ── brand colors ──
BROWN = (92, 64, 51)
GRAY = (153, 153, 153)
RED = (199, 92, 92)
GREEN_T = (107, 158, 107)

# ════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para('PRESSCO21', bold=True, size=28, color=BROWN, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('브랜드 디자인 가이드 & 토큰 활용 매뉴얼', bold=True, size=18, color=BROWN, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('꽃으로 노는 모든 방법', size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('"이 문서에 정의된 토큰(컬러, 폰트, 간격)을 따르면', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('누가 만들어도, AI가 만들어도 PRESSCO21답습니다."', size=10, color=BROWN, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
add_para('작성일: 2026-04-22', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('대상: 디자인기획팀 (장다경 팀장, 조승해 사원)', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Figma: figma.com/design/tw2xLLJ7mSLdUgZKnbYwGJ', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ════════════════════════════════════════════
# 0. 이 문서를 사용하는 법
# ════════════════════════════════════════════
add_heading('0. 이 문서를 사용하는 법', level=1)

add_para('이 가이드는 단순한 "컬러표"가 아닙니다. 아래 3가지 상황에서 펴보세요:', size=10)
doc.add_paragraph()
add_para('상황 1: ChatGPT에서 이미지를 만들 때', bold=True, size=11, color=BROWN)
add_para('  → 6장 "토큰 기반 프롬프트 작성법" + 7장 "프롬프트 템플릿"을 열고, 토큰 값을 프롬프트에 그대로 복사합니다.')
doc.add_paragraph()
add_para('상황 2: 생성된 이미지를 검수할 때', bold=True, size=11, color=BROWN)
add_para('  → 8장 "토큰 기반 QC 체크리스트"로 컬러, 톤, 텍스트, 여백을 하나씩 대조합니다.')
doc.add_paragraph()
add_para('상황 3: Figma에서 직접 디자인할 때', bold=True, size=11, color=BROWN)
add_para('  → Figma 변수 패널에서 토큰을 선택하면 자동으로 브랜드에 맞는 값이 적용됩니다.')
doc.add_paragraph()
add_para('핵심 원칙: "감"으로 색상을 고르지 말고, 이 문서의 토큰 이름으로 선택하세요.', bold=True, size=10, color=BROWN)

doc.add_page_break()

# ════════════════════════════════════════════
# 1. 브랜드 정체성
# ════════════════════════════════════════════
add_heading('1. 브랜드 정체성', level=1)
add_table(['항목', '내용'], [
    ['슬로건', '꽃으로 노는 모든 방법'],
    ['브랜드 방향', '"재료상" → "꽃 공예 전문 기업"'],
    ['톤앤무드', '우아 + 따뜻 + 자연 + 포용 + 신뢰'],
    ['고객 1순위', '수강생 (파트너클래스)'],
    ['핵심 가치', '30년 전통의 전문성 + 누구나 쉽게 시작하는 접근성'],
], col_widths=[4, 12])

# ════════════════════════════════════════════
# 2. 컬러 토큰
# ════════════════════════════════════════════
add_heading('2. 컬러 토큰', level=1)
add_para('모든 디자인은 아래 17개 컬러 안에서만 선택합니다. "이 색 예쁜데?" 가 아니라 "이건 Primary/Green이야"로 말합니다.', size=10)

add_heading('2-1. Primary — 반드시 1개 이상 사용', level=2)
add_color_table([
    {'name': 'Soft Pink', 'hex': '#F8E8E0', 'usage': '배경, 카드, 섹션 구분'},
    {'name': 'Ivory', 'hex': '#FFF8F0', 'usage': '메인 배경, 상세페이지 바탕'},
    {'name': 'Green', 'hex': '#8BA888', 'usage': '포인트, CTA 버튼, 강조'},
    {'name': 'Deep Brown', 'hex': '#5C4033', 'usage': '제목, 핵심 텍스트, 로고'},
])

add_heading('2-2. Accent — 시즌/이벤트 선택 사용', level=2)
add_color_table([
    {'name': 'Gold', 'hex': '#C9A96E', 'usage': '프리미엄/수료증/감사 (어버이날)'},
    {'name': 'Rose', 'hex': '#D4907E', 'usage': '봄/여성/로맨틱 시즌'},
    {'name': 'Sage', 'hex': '#A8BFA0', 'usage': '자연/힐링/여름 콘텐츠'},
    {'name': 'Lavender', 'hex': '#B8A9C9', 'usage': '특별 이벤트/한정판'},
])

add_heading('2-3. Neutral — 텍스트, 배경, 구분선', level=2)
add_color_table([
    {'name': 'White', 'hex': '#FFFFFF', 'usage': '기본 배경'},
    {'name': 'Light Gray', 'hex': '#F5F5F5', 'usage': '섹션 구분 배경'},
    {'name': 'Gray', 'hex': '#999999', 'usage': '보조 텍스트, 캡션'},
    {'name': 'Dark', 'hex': '#333333', 'usage': '본문 텍스트'},
    {'name': 'Black', 'hex': '#1A1A1A', 'usage': '강조 제목 (제한적)'},
])

add_heading('2-4. Semantic — UI 상태 표시', level=2)
add_color_table([
    {'name': 'Success', 'hex': '#6B9E6B', 'usage': '완료, 재고 있음'},
    {'name': 'Warning', 'hex': '#E0A84B', 'usage': '주의, 품절 임박'},
    {'name': 'Error', 'hex': '#C75C5C', 'usage': '오류, 품절'},
    {'name': 'Info', 'hex': '#5C8DC7', 'usage': '안내, 도움말'},
])

doc.add_page_break()

# ════════════════════════════════════════════
# 3. 컬러 토큰 활용 레시피
# ════════════════════════════════════════════
add_heading('3. 컬러 토큰 활용 레시피', level=1)
add_para('컬러를 "어떤 상황에서 어떻게 조합하는가"가 핵심입니다. 아래 레시피대로 따르면 됩니다.', size=10)

add_heading('3-1. 디자인 유형별 컬러 매핑', level=2)
add_table(
    ['디자인 유형', '배경', '제목', '본문', '포인트/CTA', '보조'],
    [
        ['상세페이지', 'Ivory', 'Deep Brown', 'Dark', 'Green', 'Soft Pink'],
        ['이벤트 배너', 'Soft Pink', 'Deep Brown', '—', 'Green or Gold', '시즌 Accent'],
        ['SNS 피드', 'Ivory or White', 'Deep Brown', '—', '시즌 Accent', '—'],
        ['YouTube 썸네일', 'Deep Brown or Green', 'White', '—', 'Gold', '—'],
        ['교육자료', 'White', 'Deep Brown', 'Dark', 'Green', 'Light Gray'],
        ['키트 구성도', 'White or Ivory', 'Deep Brown', 'Dark', 'Green (번호)', 'Gray (라벨)'],
    ],
    col_widths=[3.2, 2.4, 2.4, 2, 3, 2.4]
)

add_heading('3-2. 시즌별 Accent 컬러 선택 가이드', level=2)
add_table(
    ['시즌/이벤트', 'Accent 컬러', '조합 예시'],
    [
        ['어버이날 (5월)', 'Gold + Rose', '배경 Soft Pink + 타이틀 Deep Brown + 포인트 Gold'],
        ['스승의날 (5월)', 'Gold + Sage', '배경 Ivory + 타이틀 Deep Brown + 포인트 Sage'],
        ['봄/벚꽃 시즌', 'Rose', '배경 Soft Pink + 포인트 Rose + 텍스트 Deep Brown'],
        ['여름/그린 시즌', 'Sage', '배경 White + 포인트 Sage + 텍스트 Dark'],
        ['가을/추석', 'Gold', '배경 Ivory + 포인트 Gold + 텍스트 Deep Brown'],
        ['크리스마스', 'Lavender + Gold', '배경 Ivory + 포인트 Lavender + 악센트 Gold'],
        ['프리미엄/VIP', 'Gold', '배경 Deep Brown + 텍스트 White + 포인트 Gold'],
        ['신규 런칭', 'Lavender', '배경 White + 포인트 Lavender + CTA Green'],
    ],
    col_widths=[3.5, 3, 9]
)

add_heading('3-3. 컬러 조합 DO & DON\'T', level=2)

add_para('DO (권장 조합)', bold=True, color=GREEN_T)
add_bullet('Ivory 배경 + Deep Brown 제목 + Green 버튼 — ', bold_prefix='기본형: ', color=GREEN_T)
add_bullet('Soft Pink 배경 + Deep Brown 제목 + Gold 포인트 — ', bold_prefix='시즌형: ', color=GREEN_T)
add_bullet('Deep Brown 배경 + White 텍스트 + Gold 악센트 — ', bold_prefix='프리미엄: ', color=GREEN_T)

add_para("DON'T (금지)", bold=True, color=RED)
add_bullet('Green + Rose 직접 인접 (대비 부족, 탁해 보임)', color=RED)
add_bullet('Accent 컬러 2개 이상 동시 사용 (산만해짐, 1개만 선택)', color=RED)
add_bullet('Black 배경 (브랜드 톤 불일치 — Deep Brown으로 대체)', color=RED)
add_bullet('형광색, 네온, 과포화 색상 — 어떤 경우에도 금지', color=RED)

doc.add_page_break()

# ════════════════════════════════════════════
# 4. 타이포그래피 토큰
# ════════════════════════════════════════════
add_heading('4. 타이포그래피 토큰', level=1)

add_table(['용도', '서체', '비고'], [
    ['웹/앱 기본', 'Noto Sans KR', 'Google Fonts 무료'],
    ['인쇄물', '본고딕 (동일 서체)', 'Adobe/로컬'],
], col_widths=[4, 5, 7])

add_heading('4-1. 텍스트 스타일 + 활용 장면', level=2)
add_table(
    ['스타일', '크기/행간', '두께', '어디에 쓰나'],
    [
        ['Display/Hero', '40/52', 'Bold', '메인 배너 한 줄 타이틀 ("꽃으로 노는 모든 방법")'],
        ['Heading/H1', '32/44', 'Bold', '상세페이지 최상단 제목, 이벤트 페이지 제목'],
        ['Heading/H2', '24/34', 'Bold', '섹션 제목 ("키트 구성", "수강 후기")'],
        ['Heading/H3', '20/28', 'Medium', '소제목, 카드 타이틀, 카테고리명'],
        ['Body/Large', '18/28', 'Regular', '상세페이지 핵심 설명, 강조 문구'],
        ['Body/Base', '16/26', 'Regular', '일반 본문 (가장 많이 사용)'],
        ['Body/Small', '14/22', 'Regular', '부가 설명, 주의사항, 스펙'],
        ['Caption', '12/18', 'Regular', '가격 단위, 날짜, 메타 정보'],
    ],
    col_widths=[3, 2.2, 2, 8.5]
)

add_heading('4-2. 배경별 텍스트 컬러 규칙', level=2)
add_para('배경이 밝으면 Deep Brown/Dark, 배경이 어두우면 White. 예외 없습니다.')
add_table(
    ['배경 토큰', '제목 컬러', '본문 컬러', '보조 컬러'],
    [
        ['White / Ivory / Soft Pink', 'Deep Brown #5C4033', 'Dark #333333', 'Gray #999999'],
        ['Green / Deep Brown (진한 배경)', 'White #FFFFFF', 'White #FFFFFF', 'Light Gray #F5F5F5'],
    ]
)

# ════════════════════════════════════════════
# 5. 간격 & 모서리 토큰
# ════════════════════════════════════════════
add_heading('5. 간격 & 모서리 토큰', level=1)

add_para('여백을 "적당히"가 아니라 아래 값 중 하나로 정확히 지정합니다.', size=10)

add_table(
    ['토큰', '값', '활용 장면'],
    [
        ['2xs', '4px', '아이콘과 텍스트 사이'],
        ['xs', '8px', '같은 그룹 내 요소 간격'],
        ['sm', '12px', '카드 내부 패딩, 태그 여백'],
        ['md (기본)', '16px', '대부분의 기본 간격 — 이것부터 시작'],
        ['lg', '24px', '섹션 안에서 블록 간 간격'],
        ['xl', '32px', '섹션과 섹션 사이'],
        ['2xl', '48px', '대형 섹션 구분, 히어로 아래 여백'],
        ['3xl', '64px', '페이지 최상단/하단 여백'],
    ],
    col_widths=[3, 2, 11]
)

add_table(
    ['토큰', '값', '활용 장면'],
    [
        ['none', '0px', '이미지, 전폭 배너 (모서리 없이)'],
        ['sm', '4px', '입력 필드, 작은 태그'],
        ['md (기본)', '8px', '카드, 버튼 — 기본값'],
        ['lg', '16px', '큰 카드, 모달, 팝업'],
        ['full', '9999px', '원형 아바타, 둥근 태그'],
    ],
    col_widths=[3, 2, 11]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 6. 토큰 기반 프롬프트 작성법 (핵심 신규)
# ════════════════════════════════════════════
add_heading('6. 토큰 기반 프롬프트 작성법', level=1)
add_para('ChatGPT Images 2.0에서 이미지를 만들 때, 토큰을 프롬프트에 직접 녹이면 일관된 결과가 나옵니다. "예쁜 색으로 해줘"가 아니라 "배경은 #FFF8F0, 텍스트는 #5C4033"으로 지시합니다.', size=10)

add_heading('6-1. 프롬프트에 토큰을 넣는 3가지 방법', level=2)

add_para('방법 1: HEX 직접 지정 (가장 정확)', bold=True, size=10, color=BROWN)
add_box(
    '배경색: 아이보리(#FFF8F0)\n'
    '텍스트 색상: 딥브라운(#5C4033)\n'
    '버튼/포인트: 그린(#8BA888)'
)

add_para('방법 2: 분위기 키워드 + 토큰명 (자연스러운 결과)', bold=True, size=10, color=BROWN)
add_box(
    '따뜻한 아이보리 배경에 소프트 핑크 톤의 섹션 구분,\n'
    '제목은 진한 브라운 계열, 전체적으로 우아하고 따뜻한 톤'
)

add_para('방법 3: 금지 조건 명시 (실수 방지)', bold=True, size=10, color=BROWN)
add_box(
    '금지: 형광색, 네온, 비현실적 꽃 색상, AI 생성 인물 얼굴,\n'
    '세일 폭탄 그래픽, 검정 배경, 차가운 블루톤'
)

add_para('실전 팁: 방법 1+3을 함께 쓰면 가장 안정적입니다. "원하는 것 + 원하지 않는 것"을 동시에 알려주세요.', bold=True, size=10, color=BROWN)

add_heading('6-2. 디자인 유형별 프롬프트 토큰 조합표', level=2)
add_para('아래 표에서 해당 유형을 찾아 토큰 값을 프롬프트에 복사+붙여넣기 하세요.', size=10)

add_table(
    ['유형', '배경 지시', '텍스트 지시', '분위기 지시', '금지 지시'],
    [
        [
            '상세페이지\n히어로',
            '배경: 아이보리\n(#FFF8F0) 또는\n순백(#FFFFFF)',
            '한글 타이틀:\n딥브라운(#5C4033)\n굵은 서체',
            '소프트 스튜디오\n라이팅, 따뜻하고\n우아한 톤',
            '형광색, 차가운\n블루톤, 강한\n그림자'
        ],
        [
            '이벤트 배너',
            '소프트 핑크\n(#F8E8E0)',
            '타이틀: 딥브라운\n서브: 다크(#333)',
            '시즌감 + 따뜻함\n+ 축제/감사 분위기',
            '저렴한 느낌,\n네온, 과도한\n그래픽 효과'
        ],
        [
            'SNS 피드',
            '아이보리 또는\n자연 배경',
            '텍스트 최소\n(있으면 딥브라운)',
            '자연광, 파스텔,\n따뜻한 색온도',
            '과보정, 비현실적\n색감, 어두운 톤'
        ],
        [
            'YouTube\n썸네일',
            '딥브라운(#5C4033)\n또는 그린(#8BA888)',
            '흰색 굵은 폰트\n+ 외곽선',
            '강렬한 대비,\n시선 집중',
            'AI 얼굴, 클릭\n베이트 스타일'
        ],
        [
            '키트 구성도',
            '순백(#FFFFFF)\n또는 밝은 우드',
            '각 구성품 옆\n한글 라벨\n딥브라운+그레이',
            '깔끔한 플랫레이\n정보 전달 중심',
            '3D 렌더링 느낌,\n과도한 그림자'
        ],
    ],
    col_widths=[2.5, 3, 3, 3, 3]
)

add_heading('6-3. 시즌 이벤트 프롬프트 컬러 레시피', level=2)
add_para('시즌별로 Primary + 어떤 Accent를 조합하는지 미리 정해두면 매번 고민하지 않습니다.', size=10)

add_table(
    ['시즌', '프롬프트에 넣을 컬러 지시 (복사용)'],
    [
        ['어버이날', '배경: 소프트 핑크(#F8E8E0), 포인트: 골드(#C9A96E), 텍스트: 딥브라운(#5C4033), 분위기: 따뜻한 감사, 카네이션/압화'],
        ['스승의날', '배경: 아이보리(#FFF8F0), 포인트: 세이지(#A8BFA0), 텍스트: 딥브라운(#5C4033), 분위기: 존경과 감사, 자연스러운'],
        ['봄 시즌', '배경: 소프트 핑크(#F8E8E0), 포인트: 로즈(#D4907E), 분위기: 벚꽃, 화사, 새로운 시작'],
        ['여름 시즌', '배경: 화이트(#FFFFFF), 포인트: 세이지(#A8BFA0), 분위기: 시원한 자연, 그린 톤, 힐링'],
        ['가을/추석', '배경: 아이보리(#FFF8F0), 포인트: 골드(#C9A96E), 분위기: 풍성함, 감사, 따뜻한 골드톤'],
        ['크리스마스', '배경: 아이보리(#FFF8F0), 포인트: 라벤더(#B8A9C9)+골드(#C9A96E), 분위기: 특별한, 겨울 따뜻함'],
    ],
    col_widths=[2.5, 13.5]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 7. 프롬프트 템플릿 (토큰 내장)
# ════════════════════════════════════════════
add_heading('7. 프롬프트 템플릿 5종 (토큰 내장, 복사해서 사용)', level=1)

add_para('공통 가드레일 — 모든 프롬프트 끝에 붙이기', bold=True, size=11, color=BROWN)
add_box(
    '[PRESSCO21 브랜드 가드레일]\n'
    '- 메인 컬러: 소프트 핑크(#F8E8E0), 아이보리(#FFF8F0), 그린(#8BA888), 딥브라운(#5C4033)\n'
    '- 톤: 우아 + 따뜻 + 자연. 소프트 라이팅, 따뜻한 색온도\n'
    '- 금지: 형광색, AI 생성 인물 얼굴, 비현실적 꽃 색상, 세일 폭탄 그래픽, 네온, 검정 배경, 차가운 블루톤',
    bg='FFF8F0'
)

templates = [
    ('T1. 상세페이지 히어로',
     'PRESSCO21 상세페이지 히어로 이미지. 2048x2048.\n'
     '제품: ________\n'
     '구도: 제품이 중앙, 주변에 재료가 자연스럽게 배치된 플랫레이\n'
     '배경: 아이보리(#FFF8F0)\n'
     '조명: 소프트 스튜디오 라이팅, 그림자 최소\n'
     '색온도: 따뜻함 (5500~6000K 느낌)\n'
     '제목 텍스트가 필요하면: 딥브라운(#5C4033) 굵은 서체\n\n'
     '[PRESSCO21 브랜드 가드레일]'),
    ('T2. 이벤트/시즌 배너 (멀티사이즈)',
     'PRESSCO21 ________ 배너 4종 동시 생성.\n'
     '사이즈: 1200x400 / 720x480 / 860x360 / 1080x1080\n'
     '배경: 소프트 핑크(#F8E8E0)\n'
     '한글 타이틀: "________" → 딥브라운(#5C4033) 굵은 서체\n'
     '서브 카피: "________" → 다크(#333333)\n'
     '포인트 컬러: ________ (시즌 Accent 1개만 선택)\n'
     '우측 하단 PRESSCO21 로고 여백 80x80px 확보\n\n'
     '[PRESSCO21 브랜드 가드레일]'),
    ('T3. SNS 피드 이미지',
     'PRESSCO21 Instagram 피드. 1080x1080.\n'
     '장면: ________ (꽃공예 작업과정 / 완성작품 / 재료 플랫레이)\n'
     '배경: 아이보리(#FFF8F0) 또는 자연 배경\n'
     '톤: 따뜻한 자연광, 파스텔, 색온도 따뜻하게\n'
     '텍스트: ________ (있으면 하단 20%에 딥브라운 1줄, 없으면 이미지만)\n\n'
     '[PRESSCO21 브랜드 가드레일]'),
    ('T4. 키트 구성도 / 인포그래픽',
     'PRESSCO21 키트 구성 인포그래픽. 2048px.\n'
     '키트: ________\n'
     '구성품: 1. ________ 2. ________ 3. ________\n'
     '레이아웃: 플랫레이 + 각 구성품 옆에 한글 라벨(딥브라운 #5C4033) + 번호(그린 #8BA888)\n'
     '배경: 순백(#FFFFFF) 또는 밝은 우드톤\n'
     '하단에 "키트 구성 안내" 타이틀(딥브라운) + PRESSCO21 표기\n'
     '연결선/화살표: 그레이(#999999)\n\n'
     '[PRESSCO21 브랜드 가드레일]'),
    ('T5. YouTube 썸네일',
     'PRESSCO21 YouTube 썸네일. 1280x720.\n'
     '주제: ________\n'
     '좌측 60%: 작품/작업 이미지 (밝고 선명하게)\n'
     '우측 40%: 한글 타이틀 "________"\n'
     '  → 굵은 흰색(#FFFFFF) + 딥브라운(#5C4033) 외곽선\n'
     '배경: ________ (딥그린 #8BA888 / 딥브라운 #5C4033 중 택1)\n'
     '포인트: 골드(#C9A96E) 악센트\n\n'
     '[PRESSCO21 브랜드 가드레일]'),
]

for title, body in templates:
    add_para(title, bold=True, size=11, color=BROWN)
    add_box(body, bg='F5F5F5')

doc.add_page_break()

# ════════════════════════════════════════════
# 8. 토큰 기반 QC 체크리스트 (신규)
# ════════════════════════════════════════════
add_heading('8. 토큰 기반 QC 체크리스트', level=1)
add_para('AI가 생성한 이미지, 또는 직접 만든 디자인을 승인하기 전에 아래를 대조합니다.', size=10)

add_heading('8-1. 컬러 체크 (가장 중요)', level=2)
add_table(
    ['체크 항목', '기준', 'OK / NG'],
    [
        ['배경색이 토큰 범위인가?', 'White / Ivory / Soft Pink / Light Gray 중 하나', ''],
        ['제목 컬러가 맞는가?', '밝은 배경 → Deep Brown, 어두운 배경 → White', ''],
        ['Accent 컬러가 1개만 사용되었는가?', '시즌 Accent 2개 이상 동시 사용 = NG', ''],
        ['형광색/네온이 없는가?', '조금이라도 있으면 NG', ''],
        ['전체 톤이 "따뜻"한가?', '차갑거나 어두운 느낌 = NG', ''],
    ],
    col_widths=[5, 7.5, 2]
)

add_heading('8-2. 텍스트 체크', level=2)
add_table(
    ['체크 항목', '기준', 'OK / NG'],
    [
        ['한글이 정상 렌더링 되었는가?', '깨진 글자, 이상한 획 = NG', ''],
        ['텍스트 크기 위계가 맞는가?', '타이틀 > 서브 > 본문 순서', ''],
        ['텍스트가 배경 위에서 읽히는가?', '대비 부족 = NG', ''],
    ],
    col_widths=[5, 7.5, 2]
)

add_heading('8-3. 이미지 품질 체크', level=2)
add_table(
    ['체크 항목', '기준', 'OK / NG'],
    [
        ['꽃/재료 색상이 자연스러운가?', '현실에 없는 색 = NG', ''],
        ['AI 생성 인물 얼굴이 없는가?', '있으면 무조건 NG', ''],
        ['로고 삽입 여백이 확보되었는가?', '우측 하단 80x80px', ''],
        ['지정 사이즈에 맞는가?', '5장 사이즈 규격표 참조', ''],
        ['조명이 소프트한가?', '강한 그림자, 역광 = NG', ''],
    ],
    col_widths=[5, 7.5, 2]
)

add_para('판정 기준:', bold=True, size=10, color=BROWN)
add_para('  - 전체 OK → 승인, 바로 사용')
add_para('  - NG 1~2개 (경미) → 수정 지시 후 재생성')
add_para('  - NG 3개 이상 → 프롬프트 근본 수정 또는 실촬영 전환')

doc.add_page_break()

# ════════════════════════════════════════════
# 9. 이미지 규격 & AI vs 실촬영
# ════════════════════════════════════════════
add_heading('9. 이미지 가이드', level=1)

add_heading('9-1. 사이즈 규격표', level=2)
add_table(
    ['용도', '사이즈 (px)', '비고'],
    [
        ['상세페이지 히어로', '1200 x 600', 'PC'],
        ['상세페이지 콘텐츠', '860 x 자유', '본문 폭'],
        ['자사몰 배너 (PC)', '1200 x 400', ''],
        ['자사몰 배너 (모바일)', '720 x 480', ''],
        ['오픈마켓 기획전', '860 x 360', '쿠팡/네이버'],
        ['Instagram 피드', '1080 x 1080', '정사각'],
        ['Instagram 세로', '1080 x 1350', ''],
        ['Instagram 스토리', '1080 x 1920', ''],
        ['YouTube 썸네일', '1280 x 720', ''],
        ['카카오 알림톡', '800 x 400', ''],
    ],
    col_widths=[5, 4, 5]
)

add_heading('9-2. AI 생성 vs 실촬영', level=2)
add_table(
    ['AI 생성 가능', '반드시 실촬영'],
    [
        ['이벤트/시즌 배너', '신제품 메인 사진'],
        ['SNS 분위기 이미지', '압화 작품 디테일 (실제 질감)'],
        ['키트 구성도 / 인포그래픽', '수업 현장 / 인물 사진'],
        ['YouTube 썸네일 배경', '제품 색상이 핵심인 이미지'],
        ['카테고리 대표 이미지', '원장님/강사님 인물'],
    ],
    col_widths=[8, 8]
)
add_para('핵심: "구매 결정 직결 = 실촬영, 분위기/정보 전달 = AI 가능"', bold=True, color=BROWN)

doc.add_page_break()

# ════════════════════════════════════════════
# 10. 작업 프로세스
# ════════════════════════════════════════════
add_heading('10. 작업 프로세스', level=1)

add_heading('10-1. AI 이미지 제작 플로우', level=2)
steps = [
    ('1단계', '승해 사원: 이 가이드의 프롬프트 템플릿(7장)에서 해당 유형 선택'),
    ('2단계', '승해 사원: 시즌 컬러 레시피(6-3)에서 Accent 컬러 확인 → 프롬프트에 삽입'),
    ('3단계', '승해 사원: ChatGPT Pro에서 이미지 3~5장 생성'),
    ('4단계', '승해 사원: QC 체크리스트(8장)로 1차 자가 검수 → 2장 선별'),
    ('5단계', '다경 팀장에게 공유 → 브랜드 톤 최종 검수'),
    ('6단계', 'OK → 다운로드 → 메이크샵 어드민 업로드 / 수정 → 피드백 반영 재생성'),
]
for num, desc in steps:
    add_para(f'{num}  {desc}', size=10)

add_para('수정은 최대 2회전. 3회 넘으면 실촬영 전환 판정.', bold=True, size=10, color=BROWN)

add_heading('10-2. Figma 디자인 작업 플로우', level=2)
steps2 = [
    ('1단계', 'Figma 브랜드 마스터 파일의 변수(Variables)를 라이브러리로 연결'),
    ('2단계', '새 파일에서 색상 선택 시 → 변수 패널 → PRESSCO21 Colors에서 토큰으로 선택'),
    ('3단계', '간격 설정 시 → 변수 아이콘 → PRESSCO21 Spacing에서 토큰으로 선택'),
    ('4단계', '텍스트 작성 시 → 등록된 텍스트 스타일 적용 (H1, H2, Body 등)'),
    ('5단계', '"감으로" 색을 만들지 말고, 반드시 토큰에서 선택하기'),
]
for num, desc in steps2:
    add_para(f'{num}  {desc}', size=10)

# ════════════════════════════════════════════
# 11. Figma 안내
# ════════════════════════════════════════════
add_heading('11. Figma 브랜드 마스터 파일', level=1)
add_para('주소: https://www.figma.com/design/tw2xLLJ7mSLdUgZKnbYwGJ', bold=True, size=10)

add_table(['항목', '등록 내용'], [
    ['컬러 변수', '17개 (Primary 4 + Accent 4 + Neutral 5 + Semantic 4)'],
    ['Spacing 변수', '8단계 (4px ~ 64px)'],
    ['Radius 변수', '5단계 (0 ~ full)'],
    ['텍스트 스타일', '8단계 (Display ~ Caption, Noto Sans KR)'],
    ['Brand Guard Rails', 'DO 3개 + DON\'T 4개'],
], col_widths=[4, 12])

add_table(['항목', '수정 가능', '비고'], [
    ['변수 값 (컬러 HEX)', '다경 팀장만', '변경 시 지호님에게 알림'],
    ['새 컬러/스타일 추가', '다경 팀장만', '시즌 컬러 등'],
    ['가이드 페이지 레이아웃', '다경 팀장 + 승해 사원', ''],
    ['별도 작업 파일', '자유', '배너/상세 등'],
], col_widths=[5, 4, 6])

# ════════════════════════════════════════════
# 12. FAQ
# ════════════════════════════════════════════
add_heading('12. 자주 묻는 질문', level=1)

faqs = [
    ('Q. ChatGPT가 HEX 색상을 정확히 안 지키면?',
     'A. 100% 일치는 어렵습니다. "비슷한 톤" 수준이면 OK. 정확한 색이 필요하면 이미지 다운로드 후 Figma/Canva에서 색상 오버레이로 보정합니다.'),
    ('Q. 한글 텍스트가 깨지면?',
     'A. Images 2.0은 한글이 거의 완벽하지만 간혹 오류가 있습니다. 텍스트 없이 이미지만 생성 후 Figma/Canva에서 텍스트를 올리세요.'),
    ('Q. 같은 시리즈 느낌으로 여러 장 만들려면?',
     'A. 첫 장 생성 후 "위와 동일한 스타일로 [다른 내용]으로 1장 더"라고 하면 일관됩니다.'),
    ('Q. Thinking 모드는 언제 켜야 하나?',
     'A. 정확성이 중요한 것(키트 구성도, 상세 히어로)에만. SNS/배너는 일반 모드로 빠르게.'),
    ('Q. 실제 제품 사진을 참고시키려면?',
     'A. ChatGPT 대화창에 제품 사진 첨부 + "이 제품을 참고해서"라고 하면 됩니다.'),
    ('Q. 프롬프트에 토큰을 안 넣으면 어떻게 되나?',
     'A. ChatGPT가 "예쁘지만 PRESSCO21답지 않은" 이미지를 만듭니다. 토큰을 넣어야 브랜드 일관성이 유지됩니다.'),
]
for q, a in faqs:
    add_para(q, bold=True, size=10, color=BROWN)
    add_para(a, size=10)

doc.add_paragraph()
add_para('이 문서는 PRESSCO21 AI 팀에서 생성했습니다.', size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('질문이나 수정 요청은 지호님에게 전달해주세요.', size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── save ──
path = os.path.expanduser('~/Desktop/PRESSCO21_브랜드_디자인_가이드.docx')
doc.save(path)
print(f'저장 완료: {path}')
