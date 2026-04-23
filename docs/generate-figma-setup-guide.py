from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'; font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
style.paragraph_format.line_spacing = 1.5

for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

BROWN = (92, 64, 51); GRAY = (130, 130, 130); GREEN = (107, 158, 107); RED = (199, 92, 92)

def sf(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = 'Malgun Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)

def h(t, lv=1, level=None):
    hd = doc.add_heading(t, level=level if level else lv)
    for r in hd.runs: sf(r)

def P(t, bold=False, size=10.5, color=None, align=None):
    pr = doc.add_paragraph(); r = pr.add_run(t)
    sf(r, size=size, bold=bold, color=color)
    if align: pr.alignment = align
    return pr

def note(label, text, color):
    pr = doc.add_paragraph()
    r1 = pr.add_run(f'{label}  '); sf(r1, bold=True, size=10, color=color)
    r2 = pr.add_run(text); sf(r2, size=10, color=color)
    pr.paragraph_format.left_indent = Cm(0.5)

def box(text, bg='F5F0EB'):
    pr = doc.add_paragraph(); r = pr.add_run(text); sf(r, size=10)
    pr.paragraph_format.left_indent = Cm(0.5)
    ppr = pr._element.get_or_add_pPr()
    ppr.append(ppr.makeelement(qn('w:shd'),{qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):bg}))

def cbg(cell, fill):
    pr = cell._element.get_or_add_tcPr()
    pr.append(pr.makeelement(qn('w:shd'),{qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):fill}))

def tbl(headers, rows, cw=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,hdr in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(hdr); sf(r, size=9.5, bold=True); cbg(c,'F5F0EB')
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val)); sf(r, size=9.5)
    if cw:
        for i,w in enumerate(cw):
            for row in t.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph()

def fill_table(title, desc, headers, rows, cw=None):
    """빈칸 채우기 형식 테이블"""
    P(title, bold=True, size=11, color=BROWN)
    if desc: P(desc, size=10, color=GRAY)
    tbl(headers, rows, cw)


# ══════════════════════════════════════
# 표지
# ══════════════════════════════════════
doc.add_paragraph(); doc.add_paragraph()
P('PRESSCO21', bold=True, size=28, color=BROWN, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
P('우리 브랜드 에셋을', bold=True, size=18, color=BROWN, align=WD_ALIGN_PARAGRAPH.CENTER)
P('AI 시스템에 연결하기 위한 입력 가이드', bold=True, size=18, color=BROWN, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
P('2026-04-22', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
P('담당: 장다경 팀장, 조승해 사원', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
P('Figma 파일: figma.com/design/tw2xLLJ7mSLdUgZKnbYwGJ', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════
# 왜 이걸 해야 하나요
# ══════════════════════════════════════
h('이걸 왜 해야 하나요?')

P('우리 회사는 지금 AI 시스템(Claude, ChatGPT)을 사용해서 '
  '배너, 상세페이지, SNS 이미지 등을 만들고 있습니다.')
doc.add_paragraph()
P('그런데 AI가 "예쁜 디자인"은 만들 수 있어도, '
  '"우리 프레스코21다운 디자인"은 아직 잘 모릅니다.', bold=True)
doc.add_paragraph()
P('AI가 우리 브랜드에 맞는 결과를 내려면, '
  '우리 브랜드의 정확한 색상, 폰트, 디자인 패턴을 Figma에 등록해야 합니다. '
  'Figma에 등록된 정보를 AI가 자동으로 읽어서 사용하기 때문입니다.')
doc.add_paragraph()
P('쉽게 말하면:', bold=True, color=BROWN)
P('다경 팀장님과 승해 사원이 알고 있는 "우리 브랜드의 정확한 모습"을')
P('Figma라는 그릇에 담아주는 작업입니다.')
P('그러면 AI가 그 그릇에서 정보를 꺼내서 일관된 디자인을 만들 수 있습니다.')
doc.add_paragraph()

P('구체적으로 AI가 Figma에서 읽어가는 것:', bold=True, color=BROWN)
tbl(['AI가 읽는 것', '어디에 있나', '뭘 알 수 있나'], [
    ['컬러 변수', 'Variables 패널', '"우리 브랜드 초록색은 #8BA888이구나"'],
    ['텍스트 스타일', 'Text Styles 패널', '"제목은 32px Bold, 본문은 16px Regular이구나"'],
    ['컴포넌트', 'Assets 패널', '"우리 버튼은 이렇게 생겼구나, 카드는 이 구조구나"'],
    ['컴포넌트 이름', '레이어 이름', '"Button/primary는 메인 버튼이구나"'],
], cw=[3, 3.5, 9])

note('핵심', '다경 팀장님이 "우리 초록색은 사실 이 색이 아니라 저 색이야"라고 '
     'Figma에서 수정하면, 그 다음부터 AI가 만드는 모든 디자인에 수정된 색이 반영됩니다.', GREEN)

doc.add_page_break()

# ══════════════════════════════════════
# 지금 어떤 상태인가요
# ══════════════════════════════════════
h('지금 어떤 상태인가요?')

P('AI가 임시로 세팅해놓은 값들이 있습니다. '
  '하지만 이건 AI가 브랜드 문서를 보고 추정한 값이라, '
  '실제 우리 브랜드와 다를 수 있습니다.', bold=True)
doc.add_paragraph()
P('디자이너가 해야 할 일은 크게 3가지입니다:')
doc.add_paragraph()

tbl(['순서', '할 일', '설명', '예상 시간'], [
    ['1', '검수하기', 'AI가 넣어둔 값이 맞는지 확인하고, 틀린 것 수정', '1시간'],
    ['2', '채우기', 'AI가 못 넣은 것 (로고, 실제 사용 패턴, 컴포넌트) 추가', '3~5시간'],
    ['3', '발행하기', '완성된 결과를 "라이브러리"로 발행해서 AI가 읽을 수 있게', '10분'],
], cw=[1.5, 2.5, 8, 2.5])

doc.add_page_break()

# ══════════════════════════════════════
# 1단계: 검수
# ══════════════════════════════════════
h('1단계: AI가 넣은 값 검수하기')

P('Figma 파일을 열고, AI가 등록한 값들이 우리 브랜드와 맞는지 확인해주세요.')
P('틀린 것이 있으면 직접 수정하면 됩니다. 수정하면 AI가 다음부터 수정된 값을 사용합니다.', size=10, color=GRAY)

h('1-1. 컬러 검수', level=2)
P('아래 표의 "AI가 넣은 값"이 우리 브랜드에 맞는지 확인하고, '
  '"실제 우리 값" 칸에 정확한 값을 적어주세요.')
P('맞으면 그대로, 다르면 수정해주세요.', size=10, color=GRAY)

tbl(['이름', '용도', 'AI가 넣은 값', '맞나요?\n(O/X)', '실제 우리 값\n(다르면 기입)'], [
    ['Soft Pink', '배경, 카드', '#F8E8E0', '', ''],
    ['Ivory', '메인 배경', '#FFF8F0', '', ''],
    ['Green', 'CTA 버튼, 강조', '#8BA888', '', ''],
    ['Deep Brown', '제목, 로고', '#5C4033', '', ''],
    ['Gold', '프리미엄', '#C9A96E', '', ''],
    ['Rose', '봄 시즌', '#D4907E', '', ''],
    ['Sage', '자연/힐링', '#A8BFA0', '', ''],
    ['Lavender', '특별 이벤트', '#B8A9C9', '', ''],
], cw=[2.5, 2.5, 3, 1.5, 4.5])

note('확인 방법', 'Figma 우측 패널 > Variables 아이콘(x= 모양) > "PRESSCO21 Colors" 클릭 > '
     '각 색상의 값을 확인합니다. 수정하려면 색상을 더블클릭하세요.', GRAY)

P('추가 질문 — 아래에 해당하는 것이 있으면 적어주세요:', bold=True, color=BROWN)
tbl(['질문', '답변 (적어주세요)'], [
    ['위 8색 외에 우리가 쓰는 색상이 더 있나요?', ''],
    ['시즌별로 특별히 쓰는 색상이 있나요? (ex: 크리스마스 빨강)', ''],
    ['절대 쓰면 안 되는 색상이 있나요?', ''],
    ['로고 색상이 위 컬러 중 하나와 정확히 같나요, 다른가요?', ''],
], cw=[8, 7.5])

doc.add_page_break()

h('1-2. 폰트 검수', level=2)
P('AI가 "Noto Sans KR"을 기본 서체로 설정했습니다. 확인해주세요.')

tbl(['질문', '답변 (적어주세요)'], [
    ['자사몰에서 실제로 사용하는 한글 서체는?', ''],
    ['영문/숫자에 별도 서체를 쓰나요? (쓴다면 서체명)', ''],
    ['인쇄물(카탈로그, 명함 등)에 쓰는 서체는?', ''],
    ['SNS 이미지에서 쓰는 서체가 따로 있나요?', ''],
], cw=[8, 7.5])

P('AI가 넣은 글자 크기도 확인해주세요:', bold=True, color=BROWN)
tbl(['용도', 'AI가 넣은 크기', '우리가 실제로\n쓰는 크기', '맞나요?'], [
    ['페이지 제목 (H1)', '32px', '', ''],
    ['섹션 제목 (H2)', '24px', '', ''],
    ['소제목 (H3)', '20px', '', ''],
    ['일반 본문', '16px', '', ''],
    ['작은 설명', '14px', '', ''],
    ['아주 작은 글씨 (가격 단위 등)', '12px', '', ''],
    ['배너 큰 타이틀', '40px', '', ''],
], cw=[4.5, 3, 3.5, 3])

doc.add_page_break()

h('1-3. 간격/여백 검수', level=2)
P('자사몰 상세페이지나 배너에서 실제로 사용하는 여백 수치를 확인해주세요.')

tbl(['어디의 여백', 'AI가 넣은 값', '우리가 실제로\n쓰는 값', '맞나요?'], [
    ['카드 내부 패딩', '16px', '', ''],
    ['섹션과 섹션 사이', '32px', '', ''],
    ['상세페이지 좌우 여백', '—', '', ''],
    ['버튼 안쪽 좌우 여백', '16px', '', ''],
    ['이미지와 텍스트 사이', '12px', '', ''],
    ['배너 안쪽 여백', '32px', '', ''],
], cw=[5, 3, 3.5, 3])

h('1-4. 모서리 둥글기 검수', level=2)
tbl(['어디의 모서리', 'AI가 넣은 값', '실제 값', '맞나요?'], [
    ['버튼 모서리', '8px', '', ''],
    ['카드 모서리', '8px', '', ''],
    ['이미지 모서리', '4px', '', ''],
    ['입력 필드 모서리', '4px', '', ''],
], cw=[5, 3, 3.5, 3])

doc.add_page_break()

# ══════════════════════════════════════
# 2단계: 채우기
# ══════════════════════════════════════
h('2단계: AI가 못 넣은 것 채우기')

P('AI는 "값"은 넣을 수 있지만, "실제 디자인 패턴"은 디자이너만 알고 있습니다. '
  '아래 항목들은 디자이너가 직접 Figma에 만들어야 합니다.', bold=True)

h('2-1. 로고 등록', level=2)
P('우리 로고 파일을 Figma에 올려주세요.')

tbl(['항목', '해야 할 일', '파일 위치/메모'], [
    ['메인 로고 (컬러)', 'SVG 또는 고해상도 PNG를 Figma에 드래그하여 배치', ''],
    ['메인 로고 (흰색)', '어두운 배경용 흰색 버전', ''],
    ['심볼 로고', '아이콘 형태의 축약 로고 (있다면)', ''],
    ['로고 최소 크기', '이 이하로 줄이면 안 되는 크기', ''],
    ['로고 주변 여백', '로고 주위에 최소 이만큼은 비워야 하는 공간', ''],
], cw=[3.5, 7, 5])

note('방법', 'Figma에서 로고를 컴포넌트(Cmd+Option+K)로 만들어두면, '
     '나중에 배너나 상세페이지에서 드래그해서 바로 쓸 수 있습니다.', GRAY)

h('2-2. 버튼 디자인 등록', level=2)
P('우리 자사몰에서 실제로 쓰는 버튼 모양을 Figma에 만들어주세요.')
P('AI가 가이드로 제안한 버튼 스펙이 있지만, '
  '실제 자사몰 버튼과 다를 수 있으니 실제 모습대로 만들어주세요.', size=10, color=GRAY)

tbl(['버튼 종류', '어디에 쓰나', '지금 자사몰에서\n실제로 어떻게 생겼나? (메모)'], [
    ['메인 버튼\n(가장 중요한 행동)', '구매하기, 수강신청,\n장바구니 담기', ''],
    ['보조 버튼\n(덜 중요한 행동)', '더보기, 자세히 보기,\n목록으로', ''],
    ['외곽선 버튼\n(약한 행동)', '취소, 닫기, 뒤로가기', ''],
    ['텍스트 링크\n(링크처럼 보이는)', '전체보기, 카테고리 이동', ''],
], cw=[3, 3.5, 9])

P('각 버튼에 대해 확인할 것:', bold=True, color=BROWN)
tbl(['항목', '메인 버튼', '보조 버튼', '외곽선 버튼'], [
    ['배경색', '', '', ''],
    ['글자색', '', '', ''],
    ['테두리', '', '', ''],
    ['모서리 둥글기', '', '', ''],
    ['글자 크기', '', '', ''],
    ['높이', '', '', ''],
], cw=[3, 4, 4, 4])

doc.add_page_break()

h('2-3. 카드 디자인 등록', level=2)
P('자사몰에서 상품을 보여주는 카드 형태를 Figma에 만들어주세요.')

tbl(['카드 종류', '어디에 쓰나', '현재 자사몰에서\n실제로 어떻게 보이나? (메모)'], [
    ['상품 카드', '상품 목록, 검색 결과,\n추천 상품', ''],
    ['키트 카드', '키트/세트 상품 소개', ''],
    ['후기 카드', '구매 후기, 수강 후기', ''],
    ['클래스 카드', '파트너클래스 강좌 목록', ''],
], cw=[3, 4, 8.5])

P('카드 안에 보통 어떤 정보가 어떤 순서로 들어가나요?', bold=True, color=BROWN)
tbl(['순서', '상품 카드에 들어가는 것', '실제 우리 카드에\n들어가는 것 (적어주세요)'], [
    ['1', '상품 이미지', ''],
    ['2', '상품명', ''],
    ['3', '간단 설명', ''],
    ['4', '가격', ''],
    ['5', '버튼 (구매하기/담기)', ''],
    ['6', '기타 (배지, 할인율 등)', ''],
], cw=[1.5, 5, 8])

doc.add_page_break()

h('2-4. 배너 레이아웃 등록', level=2)
P('우리가 자주 만드는 배너의 기본 레이아웃을 Figma에 프레임으로 만들어주세요.')
P('각 사이즈별로 빈 프레임을 만들고, 안에 기본 구역(이미지 자리, 텍스트 자리, 로고 자리)을 잡아주세요.', size=10, color=GRAY)

tbl(['배너 이름', '사이즈', '어디에 쓰나', '우리가 보통 쓰는\n레이아웃 (메모)'], [
    ['PC 메인 배너', '1200 x 400', '자사몰 메인 PC', ''],
    ['모바일 메인 배너', '720 x 480', '자사몰 메인 모바일', ''],
    ['오픈마켓 배너', '860 x 360', '쿠팡/네이버', ''],
    ['인스타 피드', '1080 x 1080', 'Instagram', ''],
    ['인스타 스토리', '1080 x 1920', 'Instagram 스토리', ''],
    ['유튜브 썸네일', '1280 x 720', 'YouTube', ''],
], cw=[3, 2.5, 3, 7])

P('배너 관련 추가 질문:', bold=True, color=BROWN)
tbl(['질문', '답변 (적어주세요)'], [
    ['로고는 배너 어디에 넣나요? (우측 하단? 좌측 상단?)', ''],
    ['텍스트는 보통 어디에 배치하나요? (좌측? 중앙? 우측?)', ''],
    ['배너에 꼭 들어가는 공통 요소가 있나요?', ''],
    ['이미지는 보통 배너의 몇 %를 차지하나요?', ''],
], cw=[8, 7.5])

doc.add_page_break()

h('2-5. 상세페이지 블록 등록', level=2)
P('우리 상세페이지는 보통 어떤 구성으로 되어 있나요? '
  '아래에서 실제로 쓰는 블록에 체크해주시고, 순서를 적어주세요.')

tbl(['블록', '설명', '쓰나요?\n(O/X)', '순서\n(1,2,3...)', '메모 (우리는 이렇게 다름)'], [
    ['히어로\n(제품+타이틀)', '맨 위에 제품 사진과\n상품명/가격', '', '', ''],
    ['제품 특징', '이 상품의 장점\n3~4가지 나열', '', '', ''],
    ['키트 구성', '키트에 뭐가\n들어있는지', '', '', ''],
    ['만드는 방법', '단계별\n제작 안내', '', '', ''],
    ['후기', '구매/수강 후기', '', '', ''],
    ['FAQ', '자주 묻는 질문', '', '', ''],
    ['구매 유도\n(CTA)', '하단에\n"지금 구매하기"', '', '', ''],
    ['(기타)\n우리만 쓰는 블록', '', '', '', ''],
    ['(기타)\n우리만 쓰는 블록', '', '', '', ''],
], cw=[2.5, 3, 1.5, 1.5, 7])

P('각 블록의 배경색도 알려주세요:', bold=True, color=BROWN)
tbl(['블록', '현재 배경색\n(흰색? 연한 핑크? 기타?)'], [
    ['히어로', ''],
    ['제품 특징', ''],
    ['키트 구성', ''],
    ['만드는 방법', ''],
    ['후기', ''],
    ['구매 유도 (CTA)', ''],
], cw=[4, 11.5])

doc.add_page_break()

h('2-6. 사진 스타일 기준', level=2)
P('AI가 이미지를 만들거나 고를 때 참고하는 기준입니다. '
  '우리 브랜드에 맞는 사진 스타일을 적어주세요.')

tbl(['항목', '우리 기준 (적어주세요)'], [
    ['선호하는 배경', '(ex: 흰색, 아이보리, 나무결, 린넨...)'],
    ['선호하는 조명', '(ex: 자연광, 스튜디오, 따뜻한 톤...)'],
    ['선호하는 구도', '(ex: 플랫레이, 45도, 정면...)'],
    ['자주 쓰는 소품', '(ex: 나무 트레이, 린넨 천, 유리병...)'],
    ['절대 쓰면 안 되는 것', '(ex: 플라스틱 소품, 차가운 톤, 형광...)'],
    ['촬영 시 항상 지키는 규칙', '(있다면)'],
], cw=[4, 11.5])

doc.add_page_break()

# ══════════════════════════════════════
# 3단계: 발행
# ══════════════════════════════════════
h('3단계: 라이브러리 발행하기')

P('1~2단계에서 검수하고 채운 내용을 "라이브러리"로 발행하면, '
  'AI 시스템이 이 정보를 읽을 수 있게 됩니다.')
doc.add_paragraph()
P('발행 방법:', bold=True, size=11, color=BROWN)
P('1. Figma 좌측 패널 > "Assets" 탭 열기')
P('2. 상단의 책 모양 아이콘 클릭')
P('3. "Publish library" 버튼 클릭')
P('4. 목록 확인 후 "Publish" 클릭')
P('끝입니다!', bold=True, color=GREEN)
doc.add_paragraph()
P('이후 수정 사항이 생기면 같은 방법으로 다시 Publish하면 됩니다.', size=10, color=GRAY)

doc.add_page_break()

# ══════════════════════════════════════
# 정리
# ══════════════════════════════════════
h('정리: 디자이너가 해주면 AI가 할 수 있게 되는 것')

tbl(['디자이너가 해주는 것', 'AI가 할 수 있게 되는 것'], [
    ['컬러 값 검수/수정', '정확한 브랜드 색상으로 배너, 상세페이지 자동 생성'],
    ['폰트/크기 검수/수정', '정확한 글자 크기로 HTML 코드 자동 작성'],
    ['버튼 디자인 등록', '상세페이지에 일관된 버튼을 자동 삽입'],
    ['카드 디자인 등록', '상품 목록을 일관된 형태로 자동 생성'],
    ['배너 프레임 등록', '사이즈 정확한 배너를 빠르게 자동 제작'],
    ['상세페이지 블록 등록', '블록 조립식으로 상세페이지 자동 생성'],
    ['사진 스타일 기준 입력', 'ChatGPT로 브랜드에 맞는 이미지 자동 생성'],
    ['라이브러리 발행', '위 모든 것이 AI 시스템에 연결됨'],
], cw=[6, 9.5])

doc.add_paragraph()
P('결과적으로, 이 작업을 해두면:', bold=True, size=11, color=BROWN)
P('"새 배너 만들어줘" 한마디에 → AI가 정확한 사이즈, 정확한 색상, 정확한 레이아웃으로')
P('우리 브랜드에 맞는 배너를 만들 수 있게 됩니다.')
doc.add_paragraph()
P('이 작업은 한 번만 해두면 계속 효과가 있습니다.', bold=True, color=BROWN)
P('나중에 브랜드가 바뀌면 Figma에서 수정하고 다시 Publish하면 됩니다.')

doc.add_paragraph()
doc.add_paragraph()
P('질문이 있으면 지호님에게 전달해주세요.', size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

path = os.path.expanduser('~/Desktop/PRESSCO21_Figma_브랜드에셋_입력가이드.docx')
doc.save(path)
print(f'저장 완료: {path}')
